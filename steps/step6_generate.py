# step6_generate.py (single-line address + Times New Roman + underline rules)
from pathlib import Path
from datetime import datetime, date
from typing import Dict, List, Optional
import argparse, csv, re, sys
from collections import defaultdict
from io import BytesIO

from docx import Document  # pip install python-docx
from pw_auto.pw_common import WORD_TEMPLATE_PATH, NOTICES_DIR

BASE_DIR = Path(__file__).resolve().parent
DEBUG_DIR = BASE_DIR / "data" / "debug"
FONT_NAME = "Times New Roman"

# -------------------- CSV helpers --------------------
def find_latest_notices_csv() -> Optional[Path]:
    if not DEBUG_DIR.exists():
        return None
    cands = sorted(DEBUG_DIR.glob("*_notices.csv"), key=lambda p: p.stat().st_mtime, reverse=True)
    return cands[0] if cands else None

def read_csv_rows(path: Path) -> List[dict]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))

def pick(d: dict, *keys: str, default: str = "") -> str:
    if not d:
        return default
    ci = {re.sub(r"\s+", " ", k or "").strip().lower(): k for k in d.keys()}
    for k in keys:
        real = ci.get(re.sub(r"\s+", " ", k or "").strip().lower())
        if real is not None:
            val = d.get(real, "")
            return val.strip() if isinstance(val, str) else (str(val) if val is not None else default)
    return default

# -------------------- text/format helpers --------------------
def slugify_name(name: str) -> str:
    t = re.sub(r"[\\/:*?\"<>|]+", "", (name or "").strip())
    t = re.sub(r"\s+", " ", t)
    return t[:80] or "Unknown"

def short_mm_dd_yy(dt: date) -> str:
    return dt.strftime("%m-%d-%y")

def parse_mm_dd_yyyy_or_today(s: str) -> date:
    try:
        return datetime.strptime(s, "%m-%d-%Y").date()
    except Exception:
        return date.today()

def to_single_line(s: str) -> str:
    # collapse any line breaks and extra spaces into one line
    s = (s or "").replace("\r\n", " ").replace("\n", " ").replace("\r", " ")
    s = re.sub(r"\s+", " ", s).strip()
    return s

# -------------------- DOCX replacement with underline + font --------------------
UNDERLINE_ALWAYS = {"TenantAddress", "AmountDue", "DateOfUnpaidRent"}
UNDERLINE_FIRST_ONLY = {"TenantName", "Day", "Month"}

def build_token_regex(keys: List[str]) -> re.Pattern:
    inner = "|".join(re.escape(k) for k in sorted(keys, key=len, reverse=True))
    return re.compile(r"{{(" + inner + r")}}")

def apply_default_font(doc: Document, font_name: str = FONT_NAME):
    # Try to set common styles to Times New Roman so non-replaced text follows
    for name in ("Normal", "Header", "Footer", "Table Normal", "Body Text", "Body Text 2", "Body Text 3"):
        try:
            doc.styles[name].font.name = font_name
        except Exception:
            pass
    for i in range(1, 10):
        try:
            doc.styles[f"Heading {i}"].font.name = font_name
        except Exception:
            pass

def replace_in_paragraph_runs(paragraph, mapping: Dict[str, str],
                              token_re: re.Pattern,
                              seen_counts: Dict[str, int]):
    if not getattr(paragraph, "runs", None):
        return
    text = "".join(run.text for run in paragraph.runs)
    if not text:
        return

    parts = []
    last = 0
    for m in token_re.finditer(text):
        if m.start() > last:
            parts.append(("TEXT", text[last:m.start()], None))
        key = m.group(1)
        val = mapping.get(key, "")
        parts.append(("TOKEN", val, key))
        last = m.end()
    if last < len(text):
        parts.append(("TEXT", text[last:], None))

    if all(kind == "TEXT" for kind, _, _ in parts):
        return

    for r in paragraph.runs:
        r.text = ""
    if not paragraph.runs:
        paragraph.add_run("")

    for kind, payload, key in parts:
        if payload is None:
            continue
        run = paragraph.add_run()
        run.text = payload
        run.font.name = FONT_NAME  # force Times New Roman for inserted text
        if kind == "TOKEN" and key:
            if key in UNDERLINE_ALWAYS:
                run.underline = True
            elif key in UNDERLINE_FIRST_ONLY:
                if seen_counts[key] == 0:
                    run.underline = True
                seen_counts[key] += 1

def replace_in_table(table, mapping: Dict[str, str],
                     token_re: re.Pattern,
                     seen_counts: Dict[str, int]):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                replace_in_paragraph_runs(p, mapping, token_re, seen_counts)
            for t in cell.tables:
                replace_in_table(t, mapping, token_re, seen_counts)

def replace_placeholders_with_underlines(doc: Document, mapping: Dict[str, str]):
    token_re = build_token_regex(list(mapping.keys()))
    seen_counts = defaultdict(int)
    for p in doc.paragraphs:
        replace_in_paragraph_runs(p, mapping, token_re, seen_counts)
    for t in doc.tables:
        replace_in_table(t, mapping, token_re, seen_counts)
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                replace_in_paragraph_runs(p, mapping, token_re, seen_counts)
            for t in section.header.tables:
                replace_in_table(t, mapping, token_re, seen_counts)
        if section.footer:
            for p in section.footer.paragraphs:
                replace_in_paragraph_runs(p, mapping, token_re, seen_counts)
            for t in section.footer.tables:
                replace_in_table(t, mapping, token_re, seen_counts)

# -------------------- main --------------------
def main():
    ap = argparse.ArgumentParser(description="Generate 3-day notices from CSV")
    ap.add_argument("--csv", help="Path to notices CSV (from step 5). Defaults to latest data/debug/*_notices.csv")
    ap.add_argument("--out-dir", help="Output root folder. Defaults to NOTICES_DIR or ./data/notices")
    args = ap.parse_args()

    in_csv = Path(args.csv).resolve() if args.csv else find_latest_notices_csv()
    if not in_csv or not Path(in_csv).exists():
        print({"ok": False, "error": f"No notices CSV found. Use --csv or create one in {DEBUG_DIR}."})
        sys.exit(1)

    tpl = Path(WORD_TEMPLATE_PATH)
    if not tpl.exists():
        print({"ok": False, "error": f"Template not found at WORD_TEMPLATE_PATH: {tpl}"})
        sys.exit(1)

    default_out = Path(NOTICES_DIR).resolve() if str(NOTICES_DIR) else (BASE_DIR / "data" / "notices").resolve()
    out_root = Path(args.out_dir).resolve() if args.out_dir else default_out
    out_root.mkdir(parents=True, exist_ok=True)

    rows = read_csv_rows(Path(in_csv))
    if not rows:
        print({"ok": False, "error": f"No rows in {in_csv}"})
        sys.exit(1)

    # Read template once to avoid file locks
    tpl_bytes = tpl.read_bytes()

    made = 0
    outputs: List[str] = []
    used_names: Dict[str, int] = {}

    for r in rows:
        # Build mapping (flatten address to single line)
        tenant_name = pick(r, "TenantName")
        tenant_addr = to_single_line(pick(r, "TenantAddress"))
        mapping = {
            "TenantName": tenant_name,
            "TenantAddress": tenant_addr,
            "AmountDue": pick(r, "AmountDue"),
            "DateOfUnpaidRent": pick(r, "DateOfUnpaidRent"),
            "PropertyOwner": pick(r, "PropertyOwner"),
            "Date": pick(r, "Date"),
            "Day": pick(r, "Day"),
            "Month": pick(r, "Month"),
            "Year": pick(r, "Year"),
            "County": pick(r, "County"),
        }

        row_date = parse_mm_dd_yyyy_or_today(mapping["Date"])
        short_date = short_mm_dd_yy(row_date)
        row_dir = out_root / short_date
        row_dir.mkdir(parents=True, exist_ok=True)

        tenant_slug = slugify_name(tenant_name)
        base_name = f"{short_date}_{tenant_slug}.docx"
        n = used_names.get(base_name, 0)
        used_names[base_name] = n + 1
        file_name = base_name if n == 0 else f"{short_date}_{tenant_slug}_{n+1}.docx"

        out_path = row_dir / file_name

        try:
            doc = Document(BytesIO(tpl_bytes))
            apply_default_font(doc, FONT_NAME)  # set document styles to Times New Roman
            replace_placeholders_with_underlines(doc, mapping)
            doc.save(out_path)
            print(f"WRITE: {out_path}")
            outputs.append(str(out_path))
            made += 1
        except Exception as e:
            print(f"ERROR saving {out_path}: {e}")

    print({
        "ok": made > 0,
        "template": str(tpl),
        "input_csv": str(in_csv),
        "output_root": str(out_root),
        "files_written": made,
        "examples": outputs[:3],
    })
    sys.exit(0 if made > 0 else 1)

if __name__ == "__main__":
    main()
