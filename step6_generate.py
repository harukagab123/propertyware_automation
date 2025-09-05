# step6_generate.py
# Generate 3-Day Notice DOCX files from the Step 5 CSV of variables.
#
# Usage:
#   python step6_generate.py                # auto-detect latest data/debug/*_notices.csv
#   python step6_generate.py path\to\csv    # optional explicit CSV path
#
# Output:
#   <NOTICES_DIR>/<MM-DD-YY>/<MM-DD-YY>_<TenantName>.docx
#
# Requires:
#   - python-docx  (pip install python-docx)
#   - WORD_TEMPLATE_PATH and NOTICES_DIR set in .env (used by pw_common.py)

from pathlib import Path
from datetime import datetime, date
from typing import Dict, List, Optional
import csv, re, sys

from docx import Document  # pip install python-docx
from pw_common import WORD_TEMPLATE_PATH, NOTICES_DIR

BASE_DIR = Path(__file__).resolve().parent
DEBUG_DIR = BASE_DIR / "data" / "debug"

# -------------------- CSV helpers --------------------

def find_latest_notices_csv() -> Optional[Path]:
    if not DEBUG_DIR.exists():
        return None
    candidates = sorted(DEBUG_DIR.glob("*_notices.csv"), key=lambda p: p.stat().st_mtime, reverse=True)
    return candidates[0] if candidates else None

def read_csv_rows(path: Path) -> List[dict]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))

def pick(d: dict, *keys: str, default: str = "") -> str:
    """Case/space-insensitive dict getter."""
    if not d:
        return default
    ci = {re.sub(r"\s+", " ", k or "").strip().lower(): k for k in d.keys()}
    for k in keys:
        real = ci.get(re.sub(r"\s+", " ", k or "").strip().lower())
        if real is not None:
            val = d.get(real, "")
            return val.strip() if isinstance(val, str) else (str(val) if val is not None else default)
    return default

# -------------------- filename + dates --------------------

def slugify_name(name: str) -> str:
    t = re.sub(r"[\\/:*?\"<>|]+", "", (name or "").strip())
    t = re.sub(r"\s+", " ", t)
    return t[:80] or "Unknown"

def short_mm_dd_yy(dt: date) -> str:
    return dt.strftime("%m-%d-%y")

def parse_mm_dd_yyyy_or_today(s: str) -> date:
    try:
        return datetime.strptime(s, "%m-%d-%Y").date()
    except Exception:
        return date.today()

# -------------------- DOCX placeholder replacement --------------------

def _multi_replace(text: str, mapping: Dict[str, str]) -> str:
    """Replace placeholders in multiple common styles."""
    if not text:
        return text
    # Bracketed styles first
    for k, v in mapping.items():
        for pat in (f"{{{{{k}}}}}", f"[[{k}]]", f"<<{k}>>", f"<{k}>", f"${{{k}}}", f"{{{k}}}"):
            if pat in text:
                text = text.replace(pat, v)
    # Whole-word fallback for plain keys
    keys_word = [k for k in mapping if k.isidentifier() or re.fullmatch(r"[A-Za-z]+", k)]
    if keys_word:
        pattern = r"\b(" + "|".join(re.escape(k) for k in sorted(keys_word, key=len, reverse=True)) + r")\b"
        text = re.sub(pattern, lambda m: mapping.get(m.group(0), m.group(0)), text)
    return text

def _replace_in_paragraph(paragraph, mapping: Dict[str, str]):
    if not getattr(paragraph, "runs", None):
        return
    text = "".join(run.text for run in paragraph.runs)
    new_text = _multi_replace(text, mapping)
    if new_text != text:
        paragraph.runs[0].text = new_text
        for r in paragraph.runs[1:]:
            r.text = ""

def _replace_in_table(table, mapping: Dict[str, str]):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _replace_in_paragraph(p, mapping)
            for t in cell.tables:
                _replace_in_table(t, mapping)

def replace_placeholders(doc: Document, mapping: Dict[str, str]):
    for p in doc.paragraphs:
        _replace_in_paragraph(p, mapping)
    for t in doc.tables:
        _replace_in_table(t, mapping)
    for section in doc.sections:
        if section.header:
            for p in section.header.paragraphs:
                _replace_in_paragraph(p, mapping)
            for t in section.header.tables:
                _replace_in_table(t, mapping)
        if section.footer:
            for p in section.footer.paragraphs:
                _replace_in_paragraph(p, mapping)
            for t in section.footer.tables:
                _replace_in_table(t, mapping)

# -------------------- main --------------------

def main():
    # resolve input CSV
    in_csv = None
    if len(sys.argv) > 1:
        in_csv = Path(sys.argv[1]).resolve()
    else:
        in_csv = find_latest_notices_csv()
    if not in_csv or not in_csv.exists():
        print({"ok": False, "error": f"No notices CSV found. Pass a path or create one in {DEBUG_DIR}."})
        sys.exit(1)

    # paths
    tpl = Path(WORD_TEMPLATE_PATH)
    if not tpl.exists():
        print({"ok": False, "error": f"Template not found at WORD_TEMPLATE_PATH: {tpl}"})
        sys.exit(1)

    rows = read_csv_rows(in_csv)
    if not rows:
        print({"ok": False, "error": f"No rows in {in_csv}"})
        sys.exit(1)

    # Use the CSV "Date" (MM-DD-YYYY) to derive MM-DD-YY folder+filename; fallback to today if missing
    example_date = parse_mm_dd_yyyy_or_today(pick(rows[0], "Date"))
    short_date = short_mm_dd_yy(example_date)

    out_root = Path(NOTICES_DIR)
    out_dir = out_root / short_date
    out_dir.mkdir(parents=True, exist_ok=True)

    made = 0
    outputs: List[str] = []

    # to avoid overwriting duplicate names
    used_names: Dict[str, int] = {}

    for r in rows:
        mapping = {
            # expected placeholders in template
            "TenantName": pick(r, "TenantName"),
            "TenantAddress": pick(r, "TenantAddress"),
            "AmountDue": pick(r, "AmountDue"),
            "DateOfUnpaidRent": pick(r, "DateOfUnpaidRent"),
            "PropertyOwner": pick(r, "PropertyOwner"),
            "Date": pick(r, "Date"),          # MM-DD-YYYY
            "Day": pick(r, "Day"),
            "Month": pick(r, "Month"),
            "Year": pick(r, "Year"),
            "County": pick(r, "County"),
        }

        # per-row date override for folder/file if present
        row_date = parse_mm_dd_yyyy_or_today(mapping["Date"])
        row_short = short_mm_dd_yy(row_date)
        # ensure output folder matches each row’s date (keeps together if all same date)
        row_dir = out_root / row_short
        if not row_dir.exists():
            row_dir.mkdir(parents=True, exist_ok=True)

        tenant_slug = slugify_name(mapping["TenantName"])
        base_name = f"{row_short}_{tenant_slug}.docx"

        # de-dup filenames
        count = used_names.get(base_name, 0)
        used_names[base_name] = count + 1
        fname = base_name if count == 0 else f"{row_short}_{tenant_slug}_{count+1}.docx"

        out_path = row_dir / fname

        # fill template
        doc = Document(tpl)
        replace_placeholders(doc, mapping)
        doc.save(out_path)

        outputs.append(str(out_path))
        made += 1

    print({
        "ok": True,
        "template": str(tpl),
        "input_csv": str(in_csv),
        "output_dir": str(out_root),
        "folder_pattern": "<NOTICES_DIR>/<MM-DD-YY>/",
        "files_written": made,
        "examples": outputs[:3],
    })

if __name__ == "__main__":
    main()
